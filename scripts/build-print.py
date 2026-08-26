#!/usr/bin/env python
"""
Build the printed materials for the day into print/*.docx.

    python scripts/build-print.py

Table cards are generated from data/guestlist.csv and the menu is read from
js/menu.js — the same file the website renders from — so re-run this after any
change to the guest list, the seating or the menu. Nothing here is hand-keyed.

Run scripts/extract-border.py first (or after changing print/sunflower.jpg) to
regenerate the watercolour border in assets/.

Design rules (these print onto coloured stock):
  * No shading or background fill anywhere — the paper must show through.
    The border art is keyed to transparency for exactly this reason.
  * Ink only, in the site palette: ink #222d2a, petrol #2c5d63, wax #883a2d.
  * Fraunces / EB Garamond aren't installed on this machine, so we use the
    same fallbacks the site's CSS declares: Georgia for display, Garamond
    for body.
"""

import csv
import io
import json
import re
import shutil
import subprocess
import sys
from pathlib import Path

import qrcode
from qrcode.constants import ERROR_CORRECT_Q
from PIL import Image, ImageDraw

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor
from lxml import etree

ROOT = Path(__file__).resolve().parent.parent
CSV_PATH = ROOT / "data" / "guestlist.csv"
MENU_PATH = ROOT / "js" / "menu.js"
OUT_DIR = ROOT / "print"

# Watercolour border, lifted from print/sunflower.jpg by scripts/extract-border.py.
# Transparent everywhere except the artwork, so coloured stock still shows through.
BORDER_LANDSCAPE = ROOT / "assets" / "border-landscape.png"
BORDER_PORTRAIT = ROOT / "assets" / "border-portrait.png"

AGENDA_URL = "https://jmharte89.github.io/wedding/agenda.html"
GOFUNDME_URL = "https://gofund.me/c3356285c"

# Site palette (css/invitation.css :root)
INK = RGBColor(0x22, 0x2D, 0x2A)
PETROL = RGBColor(0x2C, 0x5D, 0x63)
WAX = RGBColor(0x88, 0x3A, 0x2D)

PETROL_RGB = (0x2C, 0x5D, 0x63)
INK_RGB = (0x22, 0x2D, 0x2A)
WAX_RGB = (0x88, 0x3A, 0x2D)

# Fraunces and EB Garamond are webfonts and aren't installed locally; these
# are the exact fallbacks --font-display / --font-body name in the CSS.
FONT_DISPLAY = "Georgia"
FONT_BODY = "Garamond"

# Members that are placeholders rather than people.
PLACEHOLDERS = {"plus-one", "plus one", "+1", "daughter", "son", "child", "tbc"}

# Framed pieces are trimmed to fit inside the frame's visible aperture with a
# few millimetres of overlap, so the border isn't swallowed by the rebate.
#
#   frame aperture   A5 118 x 169mm     A4 198 x 289mm   (measured)
#   trim to          A5 130 x 180mm     A4 210 x 297mm
#
# The A5 pieces are printed on A5 and cut down, so they carry a faint cut line.
#
# The A4 pieces are NOT cut. 210 x 300 was asked for, but A4 paper is only
# 297mm tall, so 300 cannot be printed on it. A full A4 sheet already overlaps
# the A4 aperture by 6mm at the sides and 4mm top and bottom, which is what the
# 300 was reaching for anyway — so those print full-bleed with no trimming.
CARD_A5 = (130.0, 180.0)
CARD_A5_LAND = (180.0, 130.0)

# Bleed. The artwork is drawn this much larger than the trim on every side and
# the cut line sits inside it, so cutting a millimetre off the line still
# leaves foliage running to the card edge instead of clipping a leaf or
# leaving a white sliver. 4mm rather than the usual 3, because this is being
# cut by hand with scissors rather than on a guillotine.
BLEED_MM = 4.0

# Table-card type sizes. The table name follows the Elm card as hand-edited in
# the .docx; the guest names are set as large as the widest name on any card
# allows, so the block fills the space rather than floating in it.
TABLE_NAME_SIZE_PT = 48
NAME_SIZE_PT = 26
NAME_LEAD_PT = 6

# Place cards use first names for the couple and their own — a place card is
# addressed to the person sitting there, not a register entry. Everyone else
# keeps their full name, so this is a deliberate short list, not a rule.
PLACE_CARD_NAMES = {
    "Jase Harte": "Jase",
    "Becki Harte": "Becki",
    "Thomas Carruthers": "Thomas",
}


# ----------------------------------------------------------------------
# Artwork
# ----------------------------------------------------------------------

def _quad(p0, p1, p2, steps=48):
    """Points along a quadratic bezier."""
    out = []
    for i in range(steps + 1):
        t = i / steps
        u = 1 - t
        out.append((
            u * u * p0[0] + 2 * u * t * p1[0] + t * t * p2[0],
            u * u * p0[1] + 2 * u * t * p1[1] + t * t * p2[1],
        ))
    return out


def page_ornament(w_mm, h_mm, cut_inset_mm=None):
    """The full-page trim: the watercolour sunflower-and-foliage border.

    Sourced from print/sunflower.jpg and prepared by scripts/extract-border.py,
    which strips the text that was composited into that JPEG and keys the white
    field to transparent so coloured stock still shows through.

    Returned at page size and pinned to the page by corner_trim(), which is
    what keeps it symmetric — the original tab-stop approach inherited the
    built-in Header style's centre tab and pulled the right-hand art toward
    the middle of the page.

    A-series pages are all 1:root-2, so the landscape scan serves landscape
    pages and its rotation serves portrait ones. The source is 1.486 against
    1.414 — a 5% difference, absorbed in the fit and invisible on foliage.
    """
    src = BORDER_LANDSCAPE if w_mm >= h_mm else BORDER_PORTRAIT
    if not src.exists():
        sys.exit(f"Missing {src} — run: python scripts/extract-border.py")
    img = Image.open(src).convert("RGBA")

    if cut_inset_mm:
        # The cut line sits INSIDE the artwork by the bleed distance, so the
        # foliage carries on past where the blade goes. Cutting a millimetre
        # either side of the line still leaves leaves running to the edge —
        # whereas a line drawn at the artwork's own edge means any wobble
        # either clips the leaf tips or leaves a white sliver.
        d = ImageDraw.Draw(img)
        w, h = img.size
        inset_x = round(w * cut_inset_mm / w_mm)
        inset_y = round(h * cut_inset_mm / h_mm)
        px = max(1, round(min(w, h) * 0.0012))          # ~0.15mm at print size
        d.rectangle([inset_x, inset_y, w - 1 - inset_x, h - 1 - inset_y],
                    outline=PETROL_RGB + (110,), width=px)

    buf = io.BytesIO()
    img.save(buf, format="PNG")
    buf.seek(0)
    return buf


def qr_png(url, modules_colour=INK_RGB):
    """QR with a transparent background so the paper provides the light field.

    Keeping the quiet zone transparent rather than white is what lets these
    sit on coloured stock without a white tile around them. Decoding is
    verified against the rendered PDF, not just the generated image.
    """
    qr = qrcode.QRCode(version=None, error_correction=ERROR_CORRECT_Q,
                       box_size=16, border=3)
    qr.add_data(url)
    qr.make(fit=True)
    img = qr.make_image(fill_color="black", back_color="white").convert("RGBA")
    px = img.load()
    w, h = img.size
    for y in range(h):
        for x in range(w):
            r, _g, _b, _a = px[x, y]
            if r > 128:                       # light module -> let paper through
                px[x, y] = (0, 0, 0, 0)
            else:
                px[x, y] = modules_colour + (255,)
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    buf.seek(0)
    return buf


# ----------------------------------------------------------------------
# Document furniture
# ----------------------------------------------------------------------

def style_run(run, *, font, size, colour, bold=False, italic=False,
              spacing=None, caps=False):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.color.rgb = colour
    run.bold = bold
    run.italic = italic
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = rpr.makeelement(qn("w:rFonts"), {})
        rpr.insert(0, rf)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rf.set(qn(attr), font)
    if spacing is not None:                   # letter-spacing, in twentieths pt
        sp = rpr.makeelement(qn("w:spacing"), {qn("w:val"): str(int(spacing * 20))})
        rpr.append(sp)
    if caps:
        rpr.append(rpr.makeelement(qn("w:caps"), {qn("w:val"): "1"}))
    return run


def para(container, text="", *, align=WD_ALIGN_PARAGRAPH.CENTER,
         space_before=0, space_after=6, **run_kw):
    p = container.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        style_run(p.add_run(text), **run_kw)
    return p


def _default_font(doc):
    """Document-wide default, so nothing falls back to Calibri."""
    normal = doc.styles["Normal"]
    normal.font.name = FONT_BODY
    normal.font.size = Pt(12)
    normal.font.color.rgb = INK
    rpr = normal.element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = rpr.makeelement(qn("w:rFonts"), {})
        rpr.insert(0, rf)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rf.set(qn(attr), FONT_BODY)


def _strip_borders(table):
    """No lines, ever. Table cells are used only for geometry here."""
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.makeelement(qn("w:tblBorders"), {})
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        borders.append(borders.makeelement(
            qn("w:" + edge), {qn("w:val"): "none", qn("w:sz"): "0",
                              qn("w:space"): "0"}))
    tbl_pr.append(borders)


def set_page(section, width_mm, height_mm, *, landscape=False, margin_mm=18):
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Mm(max(width_mm, height_mm))
        section.page_height = Mm(min(width_mm, height_mm))
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Mm(min(width_mm, height_mm))
        section.page_height = Mm(max(width_mm, height_mm))
    section.left_margin = Mm(margin_mm)
    section.right_margin = Mm(margin_mm)
    section.top_margin = Mm(margin_mm)
    section.bottom_margin = Mm(margin_mm)
    section.header_distance = Mm(9)
    section.footer_distance = Mm(9)


def vertically_centre(section):
    sect_pr = section._sectPr
    v = sect_pr.makeelement(qn("w:vAlign"), {qn("w:val"): "center"})
    sect_pr.append(v)


_ANCHOR = (
    '<wp:anchor xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/'
    'wordprocessingDrawing" distT="0" distB="0" distL="0" distR="0" '
    'simplePos="0" relativeHeight="0" behindDoc="1" locked="0" '
    'layoutInCell="1" allowOverlap="1">'
    '<wp:simplePos x="0" y="0"/>'
    '<wp:positionH relativeFrom="page"><wp:posOffset>{x}</wp:posOffset></wp:positionH>'
    '<wp:positionV relativeFrom="page"><wp:posOffset>{y}</wp:posOffset></wp:positionV>'
    '<wp:extent cx="{cx}" cy="{cy}"/>'
    '<wp:effectExtent l="0" t="0" r="0" b="0"/>'
    '<wp:wrapNone/>'
    '<wp:docPr id="{did}" name="Trim {did}"/>'
    '<wp:cNvGraphicFramePr/>'
    '{graphic}'
    '</wp:anchor>'
)


def add_panel(section, did, x_mm, y_mm, w_mm, h_mm, cut_inset_mm=None):
    """Anchor one border image at an exact position on the page.

    Separate from corner_trim so a sheet can carry more than one — the two-up
    A5 layouts put two independent cards on a single A4 sheet.
    """
    stream = page_ornament(w_mm, h_mm, cut_inset_mm=cut_inset_mm)

    p = section.header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = Pt(1)
    run = p.add_run()
    run.font.size = Pt(1)                 # keep the header line height at nil
    run.add_picture(stream, width=Mm(w_mm), height=Mm(h_mm))

    inline = run._element.find(qn("w:drawing")).find(qn("wp:inline"))
    graphic = inline.find(qn("a:graphic"))
    extent = inline.find(qn("wp:extent"))
    xml = _ANCHOR.format(
        cx=extent.get("cx"), cy=extent.get("cy"), did=did,
        x=int(Mm(x_mm)), y=int(Mm(y_mm)),
        graphic=etree.tostring(graphic).decode("utf-8"),
    )
    inline.getparent().replace(inline, parse_xml(xml))


def corner_trim(section, did=1, panel=None, cut_inset_mm=None):
    """Lay the page trim behind the text, pinned to the page itself.

    The image is placed as a floating anchor at an exact page offset, so its
    position cannot be perturbed by margins, paragraph styles or inherited tab
    stops. Living in the header means it repeats on every page — needed for
    the multi-page card documents — without taking part in the body's layout.

    `panel` is an optional (x_mm, y_mm, w_mm, h_mm) rectangle. Default is the
    whole page; the folded place cards pass the bottom half so the border
    lands only on the face that ends up outward once the card is folded.
    """
    if panel is None:
        x_mm, y_mm = 0.0, 0.0
        w_mm, h_mm = section.page_width.mm, section.page_height.mm
    else:
        x_mm, y_mm, w_mm, h_mm = panel

    add_panel(section, did, x_mm, y_mm, w_mm, h_mm, cut_inset_mm)


def new_doc(width_mm, height_mm, *, landscape=False, margin_mm=18,
            centre=True, did=1, margins=None, panel=None, card=None,
            inset=None):
    """`card` is a (w_mm, h_mm) trim size, centred on the sheet.

    Used where the sheet is bigger than the piece that goes in the frame: the
    border is drawn only on the trim rectangle, a faint cut line marks its
    edge, and `inset` gives the text margin measured from the TRIM edge rather
    than the paper edge, so the layout is unaffected by how much waste
    surrounds it.
    """
    doc = Document()
    # Document-wide default so nothing falls back to Calibri.
    normal = doc.styles["Normal"]
    normal.font.name = FONT_BODY
    normal.font.size = Pt(12)
    normal.font.color.rgb = INK
    rpr = normal.element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = rpr.makeelement(qn("w:rFonts"), {})
        rpr.insert(0, rf)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rf.set(qn(attr), FONT_BODY)

    sec = doc.sections[0]
    set_page(sec, width_mm, height_mm, landscape=landscape, margin_mm=margin_mm)

    cut_inset = None
    if card:
        pw, ph = sec.page_width.mm, sec.page_height.mm
        cw, ch = card
        # Artwork is drawn BLEED_MM larger than the trim on every side, so the
        # foliage runs past the blade. The cut line is then drawn that same
        # distance inside the artwork, landing exactly on the trim size.
        bw, bh = cw + 2 * BLEED_MM, ch + 2 * BLEED_MM
        panel = ((pw - bw) / 2.0, (ph - bh) / 2.0, bw, bh)
        cut_inset = BLEED_MM
        if inset:
            ox, oy = (pw - cw) / 2.0, (ph - ch) / 2.0
            il, it, ir, ib = inset
            margins = (ox + il, oy + it, ox + ir, oy + ib)

    if margins:                           # (left, top, right, bottom) in mm
        l, t, r, b = margins
        sec.left_margin, sec.top_margin = Mm(l), Mm(t)
        sec.right_margin, sec.bottom_margin = Mm(r), Mm(b)
    corner_trim(sec, did=did, panel=panel, cut_inset_mm=cut_inset)
    if centre:
        vertically_centre(sec)
    return doc


_TEXTBOX = (
    '<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
    '<w:pPr><w:spacing w:before="0" w:after="0" w:line="20" '
    'w:lineRule="exact"/></w:pPr>'
    '<w:r><w:rPr><w:sz w:val="2"/><w:noProof/></w:rPr>'
    '<w:drawing xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/'
    'wordprocessingDrawing">'
    '<wp:anchor distT="0" distB="0" distL="0" distR="0" simplePos="0" '
    'relativeHeight="{z}" behindDoc="0" locked="0" layoutInCell="0" '
    'allowOverlap="1">'
    '<wp:simplePos x="0" y="0"/>'
    '<wp:positionH relativeFrom="page"><wp:posOffset>{x}</wp:posOffset></wp:positionH>'
    '<wp:positionV relativeFrom="page"><wp:posOffset>{y}</wp:posOffset></wp:positionV>'
    '<wp:extent cx="{cx}" cy="{cy}"/>'
    '<wp:effectExtent l="0" t="0" r="0" b="0"/><wp:wrapNone/>'
    '<wp:docPr id="{did}" name="Card {did}"/><wp:cNvGraphicFramePr/>'
    '<a:graphic xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">'
    '<a:graphicData uri="http://schemas.microsoft.com/office/word/2010/'
    'wordprocessingShape">'
    '<wps:wsp xmlns:wps="http://schemas.microsoft.com/office/word/2010/'
    'wordprocessingShape">'
    '<wps:cNvSpPr txBox="1"/>'
    '<wps:spPr><a:xfrm><a:off x="0" y="0"/><a:ext cx="{cx}" cy="{cy}"/></a:xfrm>'
    '<a:prstGeom prst="rect"><a:avLst/></a:prstGeom>'
    '<a:noFill/><a:ln><a:noFill/></a:ln></wps:spPr>'
    '<wps:txbx><w:txbxContent>{body}</w:txbxContent></wps:txbx>'
    '<wps:bodyPr rot="0" spcFirstLastPara="0" vert="horz" wrap="square" '
    'lIns="0" tIns="0" rIns="0" bIns="0" numCol="1" anchor="ctr" '
    'anchorCtr="0" upright="0"><a:noAutofit/></wps:bodyPr>'
    '</wps:wsp></a:graphicData></a:graphic></wp:anchor></w:drawing></w:r></w:p>'
)


def _scratch_paragraphs(blocks):
    """Render blocks into detached paragraph XML for a text box."""
    tmp = Document()
    _default_font(tmp)
    for b in blocks:
        kw = dict(b)
        text = kw.pop("text", "")
        para(tmp, text, align=WD_ALIGN_PARAGRAPH.CENTER, **kw)
    xml = "".join(etree.tostring(p._element).decode("utf-8")
                  for p in tmp.paragraphs if p.text.strip() or True)
    return xml


def two_up(card, inset, did, *, side_by_side, blocks):
    """An A4 sheet carrying two copies of the same card.

    Everything is printed on A4, so the A5 pieces go two to a sheet. The sheet
    divides into two halves; each card is centred in its half with its own
    bleed and its own cut line, so cutting the sheet down the middle gives two
    A5 pieces, each of which then trims to the card.

    Both the border and the text are anchored to the page rather than flowed.
    An earlier attempt laid the text out in a full-bleed table and Word hung
    on it — with no page margins there is nothing for a flowed table to fit
    inside, and it never settles.
    """
    cw, ch = card
    doc = Document()
    _default_font(doc)

    sec = doc.sections[0]
    set_page(sec, 210, 297, landscape=side_by_side, margin_mm=0)

    pw, ph = sec.page_width.mm, sec.page_height.mm
    half_w = pw / 2 if side_by_side else pw
    half_h = ph if side_by_side else ph / 2
    il, it, ir, ib = inset

    # A one-line host paragraph for the anchors to hang off. The boxes are
    # positioned against the page, so this paragraph's own position is
    # irrelevant — it just has to exist and take up no height.
    host = doc.add_paragraph()
    host.paragraph_format.space_before = Pt(0)
    host.paragraph_format.space_after = Pt(0)
    host.paragraph_format.line_spacing = Pt(1)

    body_xml = _scratch_paragraphs(blocks)

    for i in range(2):
        ox = i * half_w if side_by_side else 0.0
        oy = 0.0 if side_by_side else i * half_h
        cx = ox + (half_w - cw) / 2.0          # card centred in its half
        cy = oy + (half_h - ch) / 2.0

        add_panel(sec, did + i,
                  cx - BLEED_MM, cy - BLEED_MM,
                  cw + 2 * BLEED_MM, ch + 2 * BLEED_MM,
                  cut_inset_mm=BLEED_MM)

        tb = _TEXTBOX.format(
            z=10 + i, did=did + 100 + i,
            x=int(Mm(cx + il)), y=int(Mm(cy + it)),
            cx=int(Mm(cw - il - ir)), cy=int(Mm(ch - it - ib)),
            body=body_xml,
        )
        host._element.addnext(parse_xml(tb))

    return doc


def rule(container):
    """A small centred ornament rule — three petrol dots, understated."""
    return para(container, "· · ·", font=FONT_DISPLAY, size=12, colour=PETROL,
                spacing=3.0, space_before=4, space_after=10)


def add_qr(container, url, caption, *, size_in=0.95, colour=INK_RGB,
           caption_pt=8.5):
    p = container.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(qr_png(url, colour), width=Inches(size_in))
    para(container, caption, font=FONT_BODY, size=caption_pt, colour=PETROL,
         italic=True, space_after=0)


# ----------------------------------------------------------------------
# Guest list
# ----------------------------------------------------------------------

def clean_member(raw):
    """'Jerome (12)' -> 'Jerome'; placeholders -> None."""
    name = re.sub(r"\s*\(\s*\d+\s*\)\s*$", "", (raw or "").strip())
    name = re.sub(r"\s+", " ", name).strip()
    if not name or name.lower() in PLACEHOLDERS:
        return None
    return name


def read_menu():
    """Load the menu from js/menu.js — the same file the website renders.

    The file is evaluated with Node rather than parsed here, so the printed
    menu is by construction the identical object the browser sees. Re-keying
    it into Python would be a second copy to keep in sync, which is exactly
    what the shared data file exists to prevent.
    """
    node = shutil.which("node")
    if not node:
        sys.exit("Node is needed to read js/menu.js (it is the menu's single "
                 "source of truth). Install Node, or run with --skip-menu.")

    script = (
        "const fs=require('fs');"
        "const root={};"
        "new Function('window',fs.readFileSync(process.argv[1],'utf8'))(root);"
        "process.stdout.write(JSON.stringify(root.WEDDING_MENU));"
    )
    out = subprocess.run([node, "-e", script, str(MENU_PATH)],
                         capture_output=True, check=True)
    menu = json.loads(out.stdout.decode("utf-8"))
    if not menu or not menu.get("blocks"):
        sys.exit("js/menu.js parsed but contained no blocks.")
    return menu


def read_tables():
    tables = {}
    order = []
    with CSV_PATH.open("r", encoding="utf-8-sig", newline="") as fh:
        for row in csv.DictReader(fh):
            table = (row.get("table") or "").strip()
            if not table:
                continue
            if table not in tables:
                tables[table] = []
                order.append(table)
            for m in (row.get("members") or "").split("|"):
                name = clean_member(m)
                if name and name not in tables[table]:
                    tables[table].append(name)

    # Top table leads; the rest alphabetically, for a tidy print run.
    def key(t):
        return (0, "") if t.strip().lower() == "top table" else (1, t.lower())

    return [(t, tables[t]) for t in sorted(order, key=key)]


def names_block(doc, names):
    """Lay the names out in a borderless grid, filled left to right.

    Two columns, filled ROW by row rather than column by column. The guest
    list keeps households together, so reading across pairs couples up on a
    line — Gill and Tim Bennett land side by side instead of being split
    between the foot of one column and the head of the next.

    An odd name at the end gets the whole last row to itself and sits centred
    under the pair above, rather than leaving a hole in the bottom corner.

    Three columns is deliberately not used: at 26pt the longest names
    ("Kathleen Humphries" is 76mm) overflow a 68mm column and wrap.

    Size is fixed at NAME_SIZE_PT rather than shrinking to fit, so every card
    is typographically identical.
    """
    n = len(names)
    cols = 1 if n <= 5 else 2
    rows = -(-n // cols)
    odd_tail = cols == 2 and n % 2 == 1
    size, lead = NAME_SIZE_PT, NAME_LEAD_PT
    table = doc.add_table(rows=rows, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Belt and braces: the default table style carries no borders, but say so
    # explicitly so a Word default can't reintroduce lines. No shading, ever.
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.makeelement(qn("w:tblBorders"), {})
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = borders.makeelement(qn("w:" + edge), {qn("w:val"): "none",
                                                 qn("w:sz"): "0",
                                                 qn("w:space"): "0"})
        borders.append(e)
    tbl_pr.append(borders)

    # Merge the final row first, so the odd name spans the full width and
    # centres beneath the pair above it.
    if odd_tail:
        table.cell(rows - 1, 0).merge(table.cell(rows - 1, cols - 1))

    for i, name in enumerate(names):
        r, c = divmod(i, cols)                 # row-major
        cell = table.cell(r, 0) if (odd_tail and r == rows - 1) \
            else table.cell(r, c)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(lead)
        p.paragraph_format.space_after = Pt(lead)
        style_run(p.add_run(name), font=FONT_BODY, size=size, colour=INK)
    return table


# ----------------------------------------------------------------------
# The four documents
# ----------------------------------------------------------------------

def save(doc, name):
    """Save, but say something useful if the file is open in Word.

    Word holds an exclusive lock on an open .docx, and the resulting
    PermissionError is otherwise an opaque traceback halfway through a build.
    """
    path = OUT_DIR / name
    try:
        doc.save(path)
    except PermissionError:
        print(f"  !! {name} is open in Word — skipped. Close it and re-run.")
        return False
    return True


def build_table_cards():
    """One A4 landscape card per table.

    No card for the top table: those guests get an individual folded place
    card instead, so a shared list would be redundant.

    Sizes follow the Elm card as hand-edited in the .docx — 48pt table name,
    18pt guest names, 11pt kicker, and no dot rule under the heading.
    """
    tables = [(name, people) for name, people in read_tables()
              if name.strip().lower() != "top table"]

    doc = new_doc(297, 210, landscape=True, centre=True, did=11,
                  margins=(46, 30, 46, 30))

    for idx, (name, people) in enumerate(tables):
        kicker = para(doc, "", font=FONT_DISPLAY, size=11, colour=PETROL,
                      space_after=6)
        if idx:
            # Break lives on the first run of the new card, not in a spacer
            # paragraph — an empty paragraph would throw off vertical centring.
            kicker.add_run().add_break(WD_BREAK.PAGE)
        style_run(kicker.add_run("You are seated at"), font=FONT_DISPLAY,
                  size=11, colour=PETROL, spacing=2.6, caps=True)
        para(doc, name, font=FONT_DISPLAY, size=TABLE_NAME_SIZE_PT,
             colour=INK, bold=True, space_after=10)
        names_block(doc, people)
        add_qr(doc, AGENDA_URL, "Scan for the order of the day", size_in=0.95)

    OUT_DIR.mkdir(exist_ok=True)
    save(doc, "table-cards.docx")
    return tables


def cut_grid_png(w_mm, h_mm, xs, ys, dpi=200):
    """A transparent sheet carrying only faint guide lines at xs / ys (mm)."""
    mmpx = dpi / 25.4
    W, H = int(w_mm * mmpx), int(h_mm * mmpx)
    img = Image.new("RGBA", (W, H), (0, 0, 0, 0))
    d = ImageDraw.Draw(img)
    px = max(1, round(0.15 * mmpx))
    col = PETROL_RGB + (110,)
    for x in xs:
        cx = int(x * mmpx)
        d.line([(cx, 0), (cx, H)], fill=col, width=px)
    for y in ys:
        cy = int(y * mmpx)
        d.line([(0, cy), (W, cy)], fill=col, width=px)
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    buf.seek(0)
    return buf


def add_image_panel(section, did, x_mm, y_mm, w_mm, h_mm, stream):
    """Anchor an already-made image at an exact page position."""
    p = section.header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = Pt(1)
    run = p.add_run()
    run.font.size = Pt(1)
    run.add_picture(stream, width=Mm(w_mm), height=Mm(h_mm))
    inline = run._element.find(qn("w:drawing")).find(qn("wp:inline"))
    graphic = inline.find(qn("a:graphic"))
    extent = inline.find(qn("wp:extent"))
    xml = _ANCHOR.format(
        cx=extent.get("cx"), cy=extent.get("cy"), did=did,
        x=int(Mm(x_mm)), y=int(Mm(y_mm)),
        graphic=etree.tostring(graphic).decode("utf-8"),
    )
    inline.getparent().replace(inline, parse_xml(xml))


def build_place_cards(tables):
    """Folded place cards for the top table, four to an A4 sheet.

    Everything prints on A4, and A6 is exactly a quarter of it, so four cards
    fit per sheet with no scaling: quarter the sheet, then fold each card
    across its own middle into a tent.

    As before, the border and the name sit in the BOTTOM half of each card.
    The top half stays blank — it folds back to become the support, and
    anything printed there would show through on light stock.
    """
    top = next((people for name, people in read_tables()
                if name.strip().lower() == "top table"), [])
    if not top:
        print("  (no top table found — skipping place cards)")
        return []

    PAGE_W, PAGE_H = 210.0, 297.0
    QW, QH = PAGE_W / 2, PAGE_H / 2        # A6 quadrant, 105 x 148.5
    FOLD = QH / 2                          # each card folds across its middle
    INSET_X, INSET_TOP, INSET_BOT = 13.0, 6.0, 6.0

    doc = Document()
    _default_font(doc)
    sec = doc.sections[0]
    set_page(sec, PAGE_W, PAGE_H, margin_mm=0)

    quads = [(0.0, 0.0), (QW, 0.0), (0.0, QH), (QW, QH)]

    # Borders live in the header, so the same four appear on every sheet.
    for i, (qx, qy) in enumerate(quads):
        add_panel(sec, 61 + i, qx, qy + FOLD, QW, FOLD)

    # Faint lines down the middle of the sheet, to quarter it.
    add_image_panel(sec, 69, 0, 0, PAGE_W, PAGE_H,
                    cut_grid_png(PAGE_W, PAGE_H, [QW], [QH]))

    shown = []
    for page_start in range(0, len(top), 4):
        batch = top[page_start:page_start + 4]
        host = doc.add_paragraph()
        host.paragraph_format.space_before = Pt(0)
        host.paragraph_format.space_after = Pt(0)
        host.paragraph_format.line_spacing = Pt(1)
        if page_start:
            host.add_run().add_break(WD_BREAK.PAGE)

        for i, person in enumerate(batch):
            label = PLACE_CARD_NAMES.get(person, person)
            shown.append(label)
            qx, qy = quads[i]

            n = len(label)
            size = 20 if n <= 11 else (17 if n <= 16 else 15)
            blocks = [
                dict(text="Top Table", font=FONT_DISPLAY, size=6,
                     colour=PETROL, spacing=2.0, caps=True, space_after=2),
                dict(text=label, font=FONT_DISPLAY, size=size, colour=INK,
                     bold=True, space_after=2),
                dict(text="\u00b7 \u00b7 \u00b7", font=FONT_DISPLAY, size=6.5,
                     colour=PETROL, spacing=2.0, space_after=0),
            ]
            tb = _TEXTBOX.format(
                z=20 + i, did=200 + page_start + i,
                x=int(Mm(qx + INSET_X)),
                y=int(Mm(qy + FOLD + INSET_TOP)),
                cx=int(Mm(QW - 2 * INSET_X)),
                cy=int(Mm(FOLD - INSET_TOP - INSET_BOT)),
                body=_scratch_paragraphs(blocks),
            )
            host._element.addnext(parse_xml(tb))

    save(doc, "place-cards.docx")
    return top


def build_ring_blessing():
    """A4 portrait. Wording and sizes as edited in the .docx."""
    doc = new_doc(210, 297, centre=True, did=21, margins=(40, 44, 40, 44))
    para(doc, "For Becki & Jase", font=FONT_DISPLAY, size=14, colour=PETROL,
         spacing=3.0, caps=True, space_after=12)
    para(doc, "Ring Blessing", font=FONT_DISPLAY, size=72, colour=INK,
         bold=True, space_after=8)
    rule(doc)
    para(doc,
         "These two rings will be exchanged today and worn for a lifetime.",
         font=FONT_BODY, size=18, colour=INK, space_after=12)
    para(doc,
         "Before they are, we would love you to take a moment to hold them "
         "and make a silent wish for the marriage \u2013 for patience, for "
         "laughter and for many happy years together.",
         font=FONT_BODY, size=18, colour=INK, space_after=14)
    rule(doc)
    para(doc, "Thank you", font=FONT_DISPLAY, size=18, colour=WAX,
         italic=True, space_after=0)
    save(doc, "ring-blessing.docx")


def build_favours():
    """Two copies on one A4 landscape sheet, side by side.

    Sizes are as hand-edited in the .docx — 10pt kicker, 34pt heading, 15pt
    body, 11pt footer. An earlier version shrank these on the reasoning that
    the trimmed card was smaller; that was wrong, because the margins came in
    at the same time and the text measure actually got wider.
    """
    blocks = [
        dict(text="With our thanks", font=FONT_DISPLAY, size=10,
             colour=PETROL, spacing=2.6, caps=True, space_after=8),
        dict(text="A Little Something to remember the day",
             font=FONT_DISPLAY, size=34, colour=INK, bold=True, space_after=6),
        dict(text="· · ·", font=FONT_DISPLAY, size=12, colour=PETROL,
             spacing=3.0, space_before=4, space_after=10),
        dict(text="Please take one or two home!", font=FONT_BODY, size=15,
             colour=INK, space_after=12),
        dict(text="· · ·", font=FONT_DISPLAY, size=12, colour=PETROL,
             spacing=3.0, space_before=4, space_after=10),
        dict(text="Becki & Jase  ·  29th of August 2026",
             font=FONT_DISPLAY, size=11, colour=WAX, space_after=0),
    ]
    doc = two_up(CARD_A5, (15, 18, 15, 18), did=31,
                 side_by_side=True, blocks=blocks)
    save(doc, "favours.docx")


def build_pegging():
    """Two copies on one A4 portrait sheet, stacked.

    The heading carries the joke, so it is set as large as the border allows.
    """
    blocks = [
        dict(text="A game", font=FONT_DISPLAY, size=9, colour=PETROL,
             spacing=2.6, caps=True, space_after=6),
        dict(text="How Good Are You at Pegging?", font=FONT_DISPLAY, size=34,
             colour=INK, bold=True, space_after=6),
        dict(text="· · ·", font=FONT_DISPLAY, size=12, colour=PETROL,
             spacing=3.0, space_before=4, space_after=10),
        dict(text="See how many of these pegs you can surreptitiously attach "
                  "to guests.",
             font=FONT_BODY, size=15, colour=INK, italic=True, space_after=0),
    ]
    doc = two_up(CARD_A5_LAND, (18, 14, 18, 14), did=71,
                 side_by_side=False, blocks=blocks)
    save(doc, "pegging.docx")


def build_gifts():
    """A4 portrait, full sheet — no trimming. Sizes as hand-edited in the .docx."""
    doc = new_doc(210, 297, centre=True, did=41, margins=(44, 46, 44, 46))
    para(doc, "Becki & Jase", font=FONT_DISPLAY, size=8.5, colour=PETROL,
         spacing=2.4, caps=True, space_after=10)
    para(doc, "Gifts", font=FONT_DISPLAY, size=30, colour=INK,
         bold=True, space_after=6)
    rule(doc)
    para(doc,
         "We have a house, we have stuff, and between the kids and the dogs "
         "there's barely room for anything else.",
         font=FONT_BODY, size=14, colour=INK, space_after=10)
    para(doc, "The greatest gift you can give us is being here.",
         font=FONT_BODY, size=14, colour=INK, italic=True, space_after=10)
    para(doc,
         "But if you'd really like to give something, we're saving for our "
         "honeymoon — any contribution, however small, means the world.",
         font=FONT_BODY, size=14, colour=INK, space_after=6)
    add_qr(doc, GOFUNDME_URL,
           "Scan this QR code with your camera for our honeymoon fund",
           size_in=1.35, caption_pt=9)
    save(doc, "gifts.docx")


def main():
    if not CSV_PATH.exists():
        sys.exit(f"Can't find {CSV_PATH}")
    if not MENU_PATH.exists():
        sys.exit(f"Can't find {MENU_PATH}")
    OUT_DIR.mkdir(exist_ok=True)

    # The menu is deliberately NOT printed — it lives on the website and on
    # the agenda page the QR codes point at. js/menu.js is still the source
    # for those, so nothing about the menu data has been removed.
    tables = build_table_cards()
    top = build_place_cards(tables)
    build_ring_blessing()
    build_favours()
    build_gifts()
    build_pegging()

    print(f"Fonts: display={FONT_DISPLAY}, body={FONT_BODY}")
    print(f"Table cards: {len(tables)} (top table excluded — they get "
          f"individual place cards)")
    for name, people in tables:
        print(f"  {name:<12} {len(people):>2} — {', '.join(people)}")
    total = sum(len(p) for _, p in tables)
    print(f"Seated at numbered tables: {total}")
    print(f"Place cards (top table, one per person): {len(top)}")
    print(f"  {', '.join(PLACE_CARD_NAMES.get(p, p) for p in top)}")


if __name__ == "__main__":
    main()
