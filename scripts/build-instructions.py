#!/usr/bin/env python
"""
Build print/printing-instructions.docx — the sheet that goes out by email
alongside the artwork.

    python scripts/build-instructions.py

Deliberately plain. This one is a working document for whoever is at the
printer, not a piece of the wedding stationery, so no border and no
decoration — just legible type and headings you can skim on a phone.

Everything prints on A4. Nothing needs A5 or A6 stock.
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / "print" / "printing-instructions.docx"

INK = RGBColor(0x22, 0x2D, 0x2A)
PETROL = RGBColor(0x2C, 0x5D, 0x63)
WAX = RGBColor(0x88, 0x3A, 0x2D)
DISPLAY = "Georgia"
BODY = "Garamond"


def _font(run, name):
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = rpr.makeelement(qn("w:rFonts"), {})
        rpr.insert(0, rf)
    for a in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rf.set(qn(a), name)


def para(doc, text="", *, font=BODY, size=11, colour=INK, bold=False,
         italic=False, align=WD_ALIGN_PARAGRAPH.LEFT, before=0, after=4):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    if text:
        r = p.add_run(text)
        r.font.size = Pt(size)
        r.font.color.rgb = colour
        r.bold = bold
        r.italic = italic
        _font(r, font)
    return p


def heading(doc, text, size=14, before=9):
    return para(doc, text, font=DISPLAY, size=size, colour=PETROL,
                bold=True, before=before, after=5)


def bullet(doc, text, bold_lead=None, size=11):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    if bold_lead:
        r = p.add_run(bold_lead)
        r.font.size = Pt(size); r.font.color.rgb = INK; r.bold = True
        _font(r, BODY)
    r = p.add_run(text)
    r.font.size = Pt(size); r.font.color.rgb = INK
    _font(r, BODY)
    return p


def build():
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = BODY
    normal.font.size = Pt(11)
    rpr = normal.element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = rpr.makeelement(qn("w:rFonts"), {})
        rpr.insert(0, rf)
    for a in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rf.set(qn(a), BODY)

    sec = doc.sections[0]
    sec.page_width, sec.page_height = Mm(210), Mm(297)
    sec.left_margin = sec.right_margin = Mm(20)
    sec.top_margin, sec.bottom_margin = Mm(15), Mm(14)

    # ---------------------------------------------------------------- head
    para(doc, "BECKI & JASE  ·  29 AUGUST 2026", font=DISPLAY, size=8,
         colour=PETROL, after=2)
    para(doc, "Printing and framing instructions", font=DISPLAY, size=18,
         colour=INK, bold=True, after=3)
    para(doc, "One file, 20 sheets, all on A4. Print the lot in one go. Some "
              "sheets carry more than one item and get cut up afterwards.",
         size=11, italic=True, colour=PETROL, after=7)

    # ------------------------------------------------------------ critical
    heading(doc, "First — print at actual size", before=4)
    para(doc, "In the print dialog, set scaling to 100%, \"Actual size\" or "
              "\"No scaling\".", size=11, after=4)
    para(doc,
         "Do NOT use \"Fit to page\", \"Shrink to fit\" or \"Scale to paper "
         "size\". Those change the dimensions by a few per cent, which is "
         "enough that the trimmed cards no longer fit the frames.",
         size=11, bold=True, colour=WAX, after=5)
    para(doc, "Print one sheet first and measure it before running the rest.",
         size=10.5, italic=True, colour=PETROL, after=2)

    # --------------------------------------------------------------- table
    heading(doc, "What to print")
    para(doc, "Everything is in one file: ALL-PRINT-MATERIALS.docx. "
              "Pages 1–10 are landscape and 11–20 portrait; the printer "
              "handles that itself, the paper goes in the same way "
              "throughout.", size=10.5, after=5)

    rows = [
        ("Pages", "What", "You get", "Then"),
        ("1–9", "Table cards", "9 table cards", "Nothing — frame as they are"),
        ("10", "Favours", "2 copies", "Cut in half, trim each"),
        ("11–17", "Place cards", "26 place cards", "Quarter each sheet, then fold"),
        ("18", "Ring blessing", "1 sign", "Nothing — frame as it is"),
        ("19", "Gifts", "1 sign", "Nothing — frame as it is"),
        ("20", "Pegging", "2 copies", "Cut in half, trim each"),
    ]
    t = doc.add_table(rows=len(rows), cols=4)
    t.style = "Table Grid"
    widths = (Mm(20), Mm(34), Mm(38), Mm(74))
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.cell(ri, ci)
            cell.width = widths[ci]
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(val)
            r.font.size = Pt(10)
            r.font.color.rgb = INK if ri else PETROL
            r.bold = ri == 0
            _font(r, BODY)

    para(doc, "All A4. Designed for coloured paper — no background fill "
              "anywhere, so the stock shows through.",
         size=10.5, italic=True, colour=PETROL, before=4, after=4)

    # ----------------------------------------------------------- the cuts
    heading(doc, "Favours and Pegging — pages 10 and 20")
    para(doc,
         "Each sheet carries two copies of the same sign. You only need one "
         "of each, so the second is a spare.",
         size=11, after=4)
    bullet(doc, "Cut the sheet in half, between the two copies.",
           "1.  ")
    bullet(doc, "On each half, cut along the faint grey line near the edge. "
                "The line comes away with the offcut, so nothing shows on the "
                "finished card.", "2.  ")
    bullet(doc, "The flowers deliberately carry on past that line, so a "
                "slightly wandering cut still has colour running to the edge. "
                "Do not try to cut around them.", "Note:  ")
    para(doc, "Finished sizes: Favours 130 × 180mm, Pegging 180 × 130mm. "
              "Anywhere within a few millimetres of the line is fine.",
         size=10.5, italic=True, colour=PETROL, before=3, after=2)
    para(doc,
         "The A5 frames only show 118 × 169mm, so a 130 × 180mm card is held "
         "by the frame instead of dropping through. The A4 frames show "
         "198 × 289mm and take a full sheet, which is why those aren't cut.",
         size=10.5, italic=True, colour=PETROL, before=2, after=2)

    # ------------------------------------------------------------ no cuts
    heading(doc, "Table cards, Ring Blessing and Gifts — pages 1–9, 18, 19")
    para(doc,
         "One item per sheet, printed edge to edge. No line, nothing to trim "
         "— straight into the frame.", size=11, after=4)

    # --------------------------------------------------------- place cards
    heading(doc, "Place cards — pages 11–17, four to a sheet")
    para(doc,
         "26 cards over 7 sheets, one per person on the top table. Not "
         "framed — each one folds into a little tent that stands on the table.",
         size=11, after=3)
    bullet(doc, "Cut each sheet into quarters, following the faint lines down "
                "the middle.", "1.  ")
    bullet(doc, "Fold each card in half across the middle, with the blank "
                "half folding backwards. Lining the corners up puts the fold "
                "in the right place.", "2.  ")
    bullet(doc, "The blank half is meant to be blank — it becomes the back "
                "support once folded.", "Note:  ")
    para(doc, "The last sheet has two cards on it, not four.",
         size=10.5, italic=True, colour=PETROL, before=3, after=0)

    OUT.parent.mkdir(exist_ok=True)
    doc.save(OUT)
    print(f"wrote {OUT}")


if __name__ == "__main__":
    build()
